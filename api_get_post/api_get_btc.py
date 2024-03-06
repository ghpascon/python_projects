import requests
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io

# URL da API que você deseja solicitar
api_url = "https://api.coindesk.com/v1/bpi/currentprice.json"

# Realiza a solicitação GET para a API
response = requests.get(api_url)

# Verifica se a solicitação foi bem sucedida (código de status 200)
if response.status_code == 200:
    # Converte a resposta para JSON
    data = response.json()
    
    # Extrai a data e hora de atualização
    time_updated = data['time']['updated']
    
    # Extrai o preço do Bitcoin
    bitcoin_price = float(data['bpi']['USD']['rate'].replace(',', ''))  # Convertendo para float e removendo a vírgula
    
    # Convertendo a string de data e hora para um objeto datetime
    time_updated = datetime.strptime(time_updated, "%b %d, %Y %H:%M:%S %Z")
    
    # Plotando o gráfico
    plt.figure(figsize=(10, 6))
    plt.plot(time_updated, bitcoin_price, marker='o', color='blue', label='Preço do Bitcoin (USD)')
    plt.title('Preço do Bitcoin ao Longo do Tempo')
    plt.xlabel('Data e Hora de Atualização')
    plt.ylabel('Preço do Bitcoin (USD)')
    plt.xticks(rotation=45)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    
    # Salvar o gráfico em um buffer de memória
    image_buffer = io.BytesIO()
    plt.savefig(image_buffer, format='png')
    image_buffer.seek(0)
    
    # Inserir o gráfico em um documento do Word
    doc = Document()
    doc.add_paragraph('Gráfico do Preço do Bitcoin ao Longo do Tempo:')
    doc.add_picture(image_buffer, width=Inches(6))
    doc.save('bitcoin_price_graph.docx')
    
    # Fechar o buffer de imagem
    image_buffer.close()
    
    plt.show()
else:
    print("Erro:", response.status_code)
