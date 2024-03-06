import requests
import matplotlib.pyplot as plt
from docx import Document
import io

def get_currency_exchange_data():
    url = "https://economia.awesomeapi.com.br/last/USD-BRL"

    try:
        response = requests.get(url)
        response.raise_for_status()  # Raise an exception for 4xx and 5xx status codes
        data = response.json()
        return data
    except requests.exceptions.RequestException as e:
        print("Error: Request failed:", e)
        return None

def create_graph(data):
    # Extrai os dados necessários
    high = float(data['USDBRL']['high'])
    low = float(data['USDBRL']['low'])
    variation = high - low

    # Cria o gráfico
    labels = ['High', 'Low', 'Variation']
    values = [high, low, variation]
    colors = ['green', 'red', 'blue']

    fig, ax = plt.subplots(figsize=(8, 6))
    ax.bar(labels, values, color=colors)

    # Adiciona os valores acima das barras
    for i, v in enumerate(values):
        ax.text(i, v + 0.1, str(round(v, 2)), ha='center', va='bottom')

    ax.set_ylabel('Value')
    ax.set_title('High, Low, and Variation')

    # Torna as bordas dos eixos invisíveis
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Salva o gráfico como uma imagem temporária
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)

    plt.close()  # Fecha a figura para liberar memória

    return img_buffer

def create_document(data, img_buffer):
    document = Document()

    # Adiciona o título
    document.add_heading('Currency Exchange Data', level=1)

    # Adiciona o gráfico
    document.add_picture(img_buffer)

    # Adiciona os dados sobre a taxa de câmbio
    for key, value in data['USDBRL'].items():
        document.add_paragraph(f'{key}: {value}')

    # Salva o arquivo .docx
    document.save('currency_exchange_data.docx')

# Obtém os dados da API
currency_data = get_currency_exchange_data()

# Cria o gráfico
if currency_data:
    img_buffer = create_graph(currency_data)
    create_document(currency_data, img_buffer)
    print("Document created successfully.")
else:
    print("Failed to retrieve currency exchange data.")
