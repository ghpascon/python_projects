import matplotlib
matplotlib.use('Agg') #roda sem GUI
from django.db import connection
from django.shortcuts import render
from .models import register_product, get_email
import os
import yfinance as yf
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta
import openpyxl
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import base64

#vars
loop=0
my_name=''
my_email=''
my_phone=''
my_pass=''
name=''
email=''
recipient = 'email'
subject = ''
message = ''
attachment_name = ''
attachment_path = ''
sender = ''
password = ''  
smtp_server = ''
smtp_port = ''

nome_documento=''
pasta='../project_ibov'
#open config file
with open('config.txt', 'r') as file:
    lines = file.readlines() 

#read each line and get the parameters
for line in lines:
    param = line.strip().split(':')
    if len(param) == 2:
        parameter, value = param
        if parameter.strip() == 'email':
            my_email = value.strip()
        elif parameter.strip() == 'pass':
            my_pass = value.strip()

#open config file
with open('config.txt', 'r') as file:
    lines = file.readlines() 


def limpar_bancos_de_dados():
    # Limpar todos os registros de ambas as tabelas
    register_product.objects.all().delete()
    get_email.objects.all().delete()



def excluir_doc():
    # Lista todos os arquivos na pasta
    arquivos = os.listdir(pasta)

    # Filtra apenas os arquivos com a extensão .docx
    arquivos_docx = [arquivo for arquivo in arquivos if arquivo.endswith(".docx")]

    # Exclui cada arquivo .docx
    for arquivo in arquivos_docx:
        caminho_arquivo = os.path.join(pasta, arquivo)
        os.remove(caminho_arquivo)

def send_email(recipient, subject, message, attachment_name, attachment_path, sender, password, smtp_server, smtp_port):
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = recipient
    msg['Subject'] = subject

    # Message body
    body = message
    msg.attach(MIMEText(body, 'plain'))

    # Attach the file
    with open(attachment_path, 'rb') as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= %s" % attachment_name)
        msg.attach(part)

    # Connect to SMTP server and send email
    try:
        server = smtplib.SMTP_SSL(smtp_server, smtp_port)
        server.login(sender, password)
        server.sendmail(sender, recipient, msg.as_string())
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print("Error sending email:", str(e))

def gerar_grafico_e_salvar(ticker, nome_arquivo, nome_grafico, data_inicio, data_fim):
    global dados
    # Convertendo as datas para o formato correto
    data_inicio = datetime.strptime(data_inicio, '%Y-%m-%d')
    data_fim = datetime.strptime(data_fim, '%Y-%m-%d')

    # Obtendo os dados do ticker
    dados = yf.download(ticker, start=data_inicio, end=data_fim)

   # Gerando o gráfico
    plt.figure(figsize=(10, 6), facecolor='white')  # Adiciona uma cor de fundo branca
    plt.plot(dados['Close'], color='blue', marker='.', linestyle='-', linewidth=2, label='Preço de Fechamento')  # Aumenta a espessura da linha e adiciona uma legenda
    plt.title(nome_grafico, fontsize=16, fontweight='bold')  # Aumenta o tamanho e o peso do título
    plt.xlabel('Data', fontsize=12)  # Aumenta o tamanho da fonte do rótulo do eixo x
    plt.ylabel('Preço (R$)', fontsize=12)  # Aumenta o tamanho da fonte do rótulo do eixo y
    plt.xticks(fontsize=10, rotation=45)  # Aumenta o tamanho dos ticks do eixo x e rotaciona os rótulos
    plt.yticks(fontsize=10)  # Aumenta o tamanho dos ticks do eixo y
    plt.grid(True, linestyle='--', alpha=0.7)  # Adiciona uma linha de grade com estilo tracejado e transparência
    plt.legend(fontsize=12)  # Ajusta o tamanho da fonte da legenda
    plt.tight_layout()  # Ajusta automaticamente a disposição dos elementos do gráfico para evitar sobreposição

    # Salvando o gráfico como uma imagem temporária
    nome_temp = nome_arquivo + '.png'
    plt.savefig(nome_temp)

    # Removendo a imagem temporária
    plt.close()
    return nome_temp
      
def gerar_docx():
    titulo = doc.add_heading('Relatório de Desempenho Financeiro', level=1)
    titulo.alignment = 1  # Alinha o título ao centro
    
    # Adicionando informações sobre o intervalo de datas ao documento
    periodo = doc.add_paragraph()
    periodo.add_run(f'Período(yyyy/mm/dd): {data_inicio} a {data_fim}')
    periodo.alignment = 1  # Alinha o período ao centro
    periodo.add_run('\n\n')

    nome_temp_ibovespa = gerar_grafico_e_salvar('^BVSP', 'ibovespa', f'Desempenho do Ibovespa de {data_inicio} a {data_fim}', data_inicio, data_fim)
    valor_max_ibovespa = dados['Close'].max()
    valor_min_ibovespa = dados['Close'].min()
    doc.add_paragraph().add_run().add_picture(nome_temp_ibovespa, width=Inches(6))
    nome_temp_dolar = gerar_grafico_e_salvar('USDBRL=X', 'dolar', f'Desempenho do Dólar de {data_inicio} a {data_fim}', data_inicio, data_fim)
    valor_max_dolar = dados['Close'].max()
    valor_min_dolar = dados['Close'].min()
    doc.add_paragraph().add_run().add_picture(nome_temp_dolar, width=Inches(6))

    # Adicione um parágrafo ao final do documento para mostrar os valores máximos e mínimos
    paragrafo_final = doc.add_paragraph()
    paragrafo_final.add_run("\n\n")
    paragrafo_final.add_run("Resumo dos Valores Máximos e Mínimos:\n").bold = True

    # Adicione os valores máximos e mínimos do gráfico do Ibovespa
    
    paragrafo_final.add_run(f"Ibovespa - Valor Máximo: {valor_max_ibovespa:.2f}, Valor Mínimo: {valor_min_ibovespa:.2f}\n")

    # Adicione os valores máximos e mínimos do gráfico do Dólar
    paragrafo_final.add_run(f"Dólar - Valor Máximo: R${valor_max_dolar:.2f}, Valor Mínimo: R${valor_min_dolar:.2f}\n")


    doc.save(nome_documento)

def home(request):
    return render(request, 'graph/home.html')

def docx_to_base64(nome_documento):
    with open(nome_documento, "rb") as arquivo:
        docx_data = arquivo.read()
        encoded_docx = base64.b64encode(docx_data)
    return encoded_docx.decode("utf-8")

def product_register(request):
    global data_inicio, data_fim, nome_documento, doc
    
    if request.method == 'POST':
        excluir_doc()
        limpar_bancos_de_dados()
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        receive_method = request.POST.get('receive_method')

        new_product = register_product.objects.create(
            start_date=start_date,
            end_date=end_date,
            receive_method=receive_method
        )
        new_product.save()

        data_inicio = start_date
        data_fim = end_date
        nome_documento = f'relatorio_{data_inicio}_{data_fim}.docx'
        doc = Document()

        gerar_docx()
        
        if receive_method == 'visualizar_online':
            with open('ibovespa.png', 'rb') as ibovespa_file:
                ibovespa_base64 = base64.b64encode(ibovespa_file.read()).decode('utf-8')
            with open('dolar.png', 'rb') as dolar_file:
                dolar_base64 = base64.b64encode(dolar_file.read()).decode('utf-8')
            return render(request, 'graph/online_view.html', {'ibovespa_base64': ibovespa_base64, 'dolar_base64': dolar_base64})
        elif receive_method == 'baixar_relatorio':
            relatorio_base64 = docx_to_base64(nome_documento)
            return render(request, 'graph/download_relatorio.html', {'relatorio_base64': relatorio_base64})
        elif receive_method == 'receber_por_email':
            return render(request, 'graph/email_data.html')


def email_get(request):
    if request.method == 'POST':
        nome_user = request.POST.get('nome_user')
        email_user = request.POST.get('email_user')
        
        new_email = get_email.objects.create(
            nome_user=nome_user,
            email_user=email_user,
            )
        new_email.save()
    
# Set up the email
    recipient = email_user
    subject = f'Relatório: {data_inicio}_{data_fim}'
    message = f'Olá {nome_user}, aqui está o relatório solicitado!!!'
    attachment_name = nome_documento
    attachment_path = nome_documento
    sender = my_email
    password = my_pass 
    smtp_server = 'smtp.gmail.com'
    smtp_port = 465
    send_email(recipient, subject, message, attachment_name, attachment_path, sender, password, smtp_server, smtp_port)

    return render(request, 'graph/email_send.html')

